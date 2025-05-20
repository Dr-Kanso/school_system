from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QLabel, QPushButton, QTableView,
    QHBoxLayout, QComboBox, QMessageBox, QGroupBox, 
    QHeaderView, QTableWidget, QTableWidgetItem, QFileDialog
)
from PySide6.QtCore import Qt, QDate
from PySide6.QtGui import QFont, QColor
from utils.firebase_client import FirebaseClient
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from docx.shared import Twips

class ReportsView(QWidget):
    """View for generating and viewing student performance reports"""
    
    def __init__(self, id_token, user_uid, is_admin=False):
        super().__init__()
        self.id_token = id_token
        self.user_uid = user_uid
        self.is_admin = is_admin
        self.firebase = FirebaseClient(id_token=self.id_token)
        self.standard_year_groups = ["Y7", "Y8", "Y9", "Y10", "Y11"]
        self.students_by_year = {}  # Cache students by year group
        self.setup_ui()
        self.load_terms()
    
    def setup_ui(self):
        """Set up the user interface"""
        main_layout = QVBoxLayout()
        
        # Title
        title_label = QLabel("Student Performance Report")
        title_font = QFont()
        title_font.setPointSize(14)
        title_font.setBold(True)
        title_label.setFont(title_font)
        main_layout.addWidget(title_label)
        
        # Report selection group
        selection_group = QGroupBox("Report Parameters")
        selection_layout = QVBoxLayout()
        
        # Report filters
        filters_layout = QHBoxLayout()
        
        # Year group filter
        filters_layout.addWidget(QLabel("Year Group:"))
        self.year_group_filter = QComboBox()
        self.year_group_filter.addItem("-- Select Year Group --")
        self.year_group_filter.addItems(self.standard_year_groups)
        self.year_group_filter.currentTextChanged.connect(self.on_year_group_changed)
        filters_layout.addWidget(self.year_group_filter)
        
        # Student filter - will be populated when year group is selected
        filters_layout.addWidget(QLabel("Student:"))
        self.student_filter = QComboBox()
        self.student_filter.addItem("-- Select Student --")
        self.student_filter.setEnabled(False)
        filters_layout.addWidget(self.student_filter)
        
        # Term filter
        filters_layout.addWidget(QLabel("Term:"))
        self.term_filter = QComboBox()
        self.term_filter.addItem("-- Select Term --")
        filters_layout.addWidget(self.term_filter)
        
        selection_layout.addLayout(filters_layout)
        
        # Generate report button
        self.generate_report_button = QPushButton("Generate Report")
        self.generate_report_button.clicked.connect(self.generate_report)
        selection_layout.addWidget(self.generate_report_button)
        
        selection_group.setLayout(selection_layout)
        main_layout.addWidget(selection_group)
        
        # Report display area
        report_display_group = QGroupBox("Performance Summary")
        report_display_layout = QVBoxLayout()
        
        # Student info header
        self.student_info_label = QLabel("")
        report_display_layout.addWidget(self.student_info_label)
        
        # Report table - update headers to match input_results_view
        self.report_table = QTableWidget()
        self.report_table.setColumnCount(6)  # Subject + 5 grade categories
        self.report_table.setHorizontalHeaderLabels([
            "Subject", "Current Grade", "Target Grade", "Homework", "Behaviour", "Punctuality"
        ])
        self.report_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        report_display_layout.addWidget(self.report_table)
        
        # Export button - change label to Export to DOCX
        self.export_button = QPushButton("Export to DOCX")
        self.export_button.clicked.connect(self.export_report)
        report_display_layout.addWidget(self.export_button)
        
        report_display_group.setLayout(report_display_layout)
        main_layout.addWidget(report_display_group)
        
        self.setLayout(main_layout)
    
    def load_terms(self):
        """Load available academic terms"""
        try:
            terms = self.firebase.get_collection("terms")
            
            self.term_filter.clear()
            self.term_filter.addItem("-- Select Term --", "")
            
            for term in terms:
                term_id = term.get('id', '')
                term_name = term.get('name', '')
                term_year = term.get('year', '')
                display_text = f"{term_name} {term_year}"
                self.term_filter.addItem(display_text, term_id)
                
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to load terms: {str(e)}")
    
    def on_year_group_changed(self, year_group):
        """Handle year group selection change"""
        if year_group == "-- Select Year Group --":
            self.student_filter.clear()
            self.student_filter.addItem("-- Select Student --")
            self.student_filter.setEnabled(False)
            return
            
        # Load students for the selected year group
        self.load_students_by_year_group(year_group)
    
    def load_students_by_year_group(self, year_group):
        """Load students for the selected year group"""
        try:
            # Check if we already have this year group's students cached
            if year_group in self.students_by_year:
                self.update_student_selector(self.students_by_year[year_group])
                return
                
            # Query students by year group
            students = self.firebase.query_collection_with_filters(
                "students", 
                [("year_group", "==", year_group)]
            )
            
            # Cache the results
            self.students_by_year[year_group] = students
            
            # Update the student selector
            self.update_student_selector(students)
            
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to load students: {str(e)}")
    
    def update_student_selector(self, students):
        """Update the student selector dropdown with the provided students"""
        self.student_filter.clear()
        self.student_filter.addItem("-- Select Student --", "")
        
        if not students:
            self.student_filter.setEnabled(False)
            return
            
        # Enable the selector and populate it
        self.student_filter.setEnabled(True)
        
        # Sort students by name for easier selection
        sorted_students = sorted(
            students, 
            key=lambda s: f"{s.get('first_name', '')} {s.get('last_name', '')}"
        )
        
        for student in sorted_students:
            student_id = student.get('id', '')
            first_name = student.get('first_name', '')
            last_name = student.get('last_name', '')
            
            # Use both first and last name if available
            if first_name or last_name:
                display_name = f"{first_name} {last_name}".strip()
            else:
                # Fallback to name field for backward compatibility
                display_name = student.get('name', 'Unknown')
                
            self.student_filter.addItem(display_name, student_id)
    
    def generate_report(self):
        """Generate student performance report"""
        # Get selected values
        year_group = self.year_group_filter.currentText()
        student_id = self.student_filter.currentData()
        term_id = self.term_filter.currentData()
        
        # Input validation
        if year_group == "-- Select Year Group --" or not student_id or not term_id:
            QMessageBox.warning(
                self, 
                "Selection Error", 
                "Please select a year group, student, and term."
            )
            return
            
        try:
            # Show loading indicator
            self.setCursor(Qt.WaitCursor)
            
            # Get the student details for display
            student_name = self.student_filter.currentText()
            term_name = self.term_filter.currentText()
            
            # Update the student info display
            self.student_info_label.setText(
                f"<b>Student:</b> {student_name} | <b>Year Group:</b> {year_group} | <b>Term:</b> {term_name}"
            )
            
            # Query Firestore for all results matching this term
            results = self.firebase.query_collection_with_filters(
                "results", 
                [("term_id", "==", term_id)]
            )
            
            # Clear existing table data
            self.report_table.setRowCount(0)
            
            # Process results
            if not results:
                QMessageBox.information(self, "No Data", f"No results found for {term_name}")
                self.setCursor(Qt.ArrowCursor)
                return
                
            # Find all subjects this student has grades for
            student_grades = []
            
            for result in results:
                # Check if this document contains results for our student
                student_results = result.get('student_results', {})
                if student_id in student_results:
                    subject = result.get('subject', 'Unknown')
                    grades_data = student_results[student_id]
                    
                    # Handle both old format (string) and new format (dict)
                    grades = {
                        "current_grade": "",
                        "target_grade": "", 
                        "homework": "", 
                        "behaviour": "", 
                        "punctuality": ""
                    }
                    
                    if isinstance(grades_data, str):
                        # Legacy format - just contains achievement/current grade
                        grades["current_grade"] = grades_data
                    elif isinstance(grades_data, dict):
                        # Handle both old and new category names
                        if "achievement" in grades_data:
                            grades["current_grade"] = grades_data["achievement"]
                        if "target" in grades_data:
                            grades["target_grade"] = grades_data["target"]
                        
                        # Handle current naming convention
                        for category in grades.keys():
                            if category in grades_data:
                                grades[category] = grades_data[category]
                                
                    student_grades.append({
                        'subject': subject,
                        'grades': grades
                    })
            
            # Sort by subject name
            student_grades.sort(key=lambda x: x['subject'])
            
            # Display results in the table
            for i, grade_info in enumerate(student_grades):
                self.report_table.insertRow(i)
                
                # Get the subject and grades
                subject = grade_info['subject']
                grades = grade_info['grades']
                
                # Set the subject cell
                self.report_table.setItem(i, 0, QTableWidgetItem(subject))
                
                # Set the grade cells for each category using updated category names
                categories = ["current_grade", "target_grade", "homework", "behaviour", "punctuality"]
                for j, category in enumerate(categories):
                    value = grades.get(category, "")
                    item = QTableWidgetItem(value)
                    item.setTextAlignment(Qt.AlignCenter)
                    
                    # Remove color formatting from the view - only keep it in the export
                    self.report_table.setItem(i, j + 1, item)
            
            # Optionally, fetch attendance data as well if we want to show percentage
            self.load_attendance_data(student_id, year_group, term_id)
            
            self.setCursor(Qt.ArrowCursor)
            
        except Exception as e:
            self.setCursor(Qt.ArrowCursor)
            QMessageBox.warning(self, "Error", f"Failed to generate report: {str(e)}")
            print(f"Error generating report: {str(e)}")

    def load_attendance_data(self, student_id, year_group, term_id):
        """Load attendance data for the student"""
        try:
            # This would need to analyze attendance records from the attendance collection
            # and calculate attendance percentages
            # For now, let's just show a placeholder message about attendance
            
            # Get the term to find its start and end dates
            term_data = self.firebase.get_document("terms", term_id)
            
            if not term_data:
                print(f"Term data not found for ID: {term_id}")
                return
                
            # In a real implementation, we would:
            # 1. Find all attendance records for this student within the term date range
            # 2. Calculate the attendance percentage
            # 3. Display the percentage in a dedicated label or table row
            
        except Exception as e:
            print(f"Error loading attendance data: {str(e)}")
    
    def export_report(self):
        """Export the current report to DOCX with enhanced styling"""
        if self.report_table.rowCount() == 0:
            QMessageBox.warning(self, "No Data", "There is no report data to export.")
            return
            
        try:
            # Get file save location from user
            file_path, _ = QFileDialog.getSaveFileName(
                self, "Save Report as DOCX", "", "Word Documents (*.docx)"
            )
            
            if not file_path:  # User canceled the dialog
                return
                
            # Add .docx extension if not present
            if not file_path.lower().endswith(".docx"):
                file_path += ".docx"
                
            # Parse the student info from the label text
            student_info = self.student_info_label.text()
            
            # Create a new Word document with custom page margins
            doc = Document()
            sections = doc.sections
            for section in sections:
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)
                section.left_margin = Cm(2.5)
                section.right_margin = Cm(2.5)
                # Set to landscape orientation for better table display
                section.orientation = WD_ORIENT.LANDSCAPE
            
            # Add school details at the top
            school_para = doc.add_paragraph()
            school_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            school_run = school_para.add_run("IRIS SCHOOL")
            school_run.bold = True
            school_run.font.size = Pt(16)
            
            # Address and contact information
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.add_run("Address: 100 Carlton Vale, London NW6 5HE\n").font.size = Pt(10)
            contact_para.add_run("Telephone: 02073728051 | ").font.size = Pt(10)
            contact_para.add_run("Email: irisschool@gmail.com").font.size = Pt(10)
            
            # Add title
            title = doc.add_heading('Student Performance Report', level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add a simple horizontal line using paragraph border
            border_para = doc.add_paragraph()
            border_para.paragraph_format.space_before = Pt(0)
            border_para.paragraph_format.space_after = Pt(12)
            
            # Add bottom border using simplified XML approach
            pPr = border_para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')  # Border size
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), '4F81BD')  # Blue line color
            pBdr.append(bottom)
            pPr.append(pBdr)
            
            # Extract student name, year group, and term from the info text
            student_parts = student_info.replace('<b>', '').replace('</b>', '').split('|')
            student_name = student_parts[0].split(':')[1].strip() if len(student_parts) > 0 else "Unknown"
            year_group = student_parts[1].split(':')[1].strip() if len(student_parts) > 1 else "Unknown"
            term_name = student_parts[2].split(':')[1].strip() if len(student_parts) > 2 else "Unknown"
            
            # Add student information
            info_para = doc.add_paragraph()
            info_para.add_run("Student: ").bold = True
            info_para.add_run(student_name)
            info_para.add_run(" | ")
            info_para.add_run("Year Group: ").bold = True
            info_para.add_run(year_group)
            info_para.add_run(" | ")
            info_para.add_run("Term: ").bold = True
            info_para.add_run(term_name)
            
            # Add space before grades table
            doc.add_paragraph()
            
            # Add a full table for performance summary with all columns
            column_headers = ["Subject", "Current Grade", "Target Grade", "Homework", "Behaviour", "Punctuality"]
            num_columns = len(column_headers)
            
            # Create the table
            table = doc.add_table(rows=1, cols=num_columns)
            table.style = 'Table Grid'
            
            # Add headers
            header_cells = table.rows[0].cells
            for i, header in enumerate(column_headers):
                header_cells[i].text = header
            
            # Make headers bold
            for cell in header_cells:
                # Add dark gray background to header cells
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D0D0D0"/>')
                cell._element.get_or_add_tcPr().append(shading_elm)
                
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(11)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add data rows with all performance metrics
            for row in range(self.report_table.rowCount()):
                cells = table.add_row().cells
                
                # Fill in the row with all columns from the table
                for col in range(min(num_columns, self.report_table.columnCount())):
                    item = self.report_table.item(row, col)
                    value = item.text() if item else ""
                    cells[col].text = value
                    
                    # Center-align all cells except Subject
                    if col > 0:
                        cells[col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Apply color coding based on grade value for the grade columns
                        if col in [1, 2, 3, 4, 5]: # Grade columns
                            try:
                                grade_value = int(value) if value else 0
                                cell_shading = OxmlElement('w:shd')
                                
                                if grade_value >= 7:  # High grades
                                    cell_shading.set(qn('w:fill'), "C6E0B4")  # Light green
                                elif grade_value <= 3:  # Low grades
                                    cell_shading.set(qn('w:fill'), "F8CBAD")  # Light red
                                    
                                cells[col]._tc.get_or_add_tcPr().append(cell_shading)
                            except (ValueError, TypeError):
                                pass
            
            # Apply alternating row colors for readability
            for i, row in enumerate(table.rows):
                if i % 2 == 1:  # Alternating rows
                    for cell in row.cells:
                        # Get current shading or create new one
                        tcPr = cell._tc.get_or_add_tcPr()
                        shading_elem = tcPr.xpath('./w:shd')
                        
                        # Only add shading if there's not already one from grade coloring
                        if not shading_elem:
                            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
                            tcPr.append(shading_elm)
            
            # Add grade scale information
            doc.add_paragraph().add_run("\n") # Add space
            doc.add_paragraph().add_run("Grade Scale Reference:").bold = True
            
            # Create a grade scale table with simplified styling
            scale_table = doc.add_table(rows=3, cols=2)
            scale_table.style = 'Light Grid'
            
            # Add scale information
            grades_info = [
                ("Grades 7-9", "High Achievement"),
                ("Grades 4-6", "Satisfactory Achievement"),
                ("Grades 1-3", "Needs Improvement")
            ]
            
            for i, (grades, description) in enumerate(grades_info):
                scale_table.cell(i, 0).text = grades
                scale_table.cell(i, 1).text = description
                
                # Center the grade text
                scale_table.cell(i, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add appropriate background color to match our grade coloring
                shading_elm = None
                if i == 0:  # High grades
                    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="C6E0B4"/>')
                elif i == 2:  # Low grades
                    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8CBAD"/>')
                    
                if shading_elm:
                    scale_table.cell(i, 1)._tc.get_or_add_tcPr().append(shading_elm)
                
                # Bold the grade column
                for run in scale_table.cell(i, 0).paragraphs[0].runs:
                    run.bold = True
                    run.font.size = Pt(9)
                
                # Regular font for description
                for run in scale_table.cell(i, 1).paragraphs[0].runs:
                    run.font.size = Pt(9)
            
            # Add date at the bottom
            doc.add_paragraph().add_run("\n") # Add space
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_run = date_para.add_run(f"Generated on {QDate.currentDate().toString('MMMM d, yyyy')}")
            date_run.italic = True
            date_run.font.size = Pt(9)
            
            # Save the document
            doc.save(file_path)
            
            QMessageBox.information(
                self, 
                "Export Successful", 
                f"Report has been exported to:\n{file_path}"
            )
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export report: {str(e)}")
            print(f"Error during export: {str(e)}")
